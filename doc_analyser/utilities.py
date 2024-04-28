import streamlit as st
import docx
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter

def get_split_texts(input_text,chunk_size = 30000,chunk_overlap = 1000):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size,
                     chunk_overlap=chunk_overlap,
                     length_function=len,
                     add_start_index=True)
    texts = text_splitter.create_documents([input_text])
    return texts



# Document Related Utilities
def read_docx_file(text, docx_file_name):
    doc = docx.Document(docx_file_name)
    # Iterating tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    # Iterating paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text
    return text


def read_pdf_file(text, pdf_file_name):
    pdf_reader = PyPDF2.PdfReader(pdf_file_name)
    for page_no in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page_no].extract_text()
    return text



def set_page_config(page_title, page_icon="🔬"):
    """ Sets the theme for the app webpage"""
    st.set_page_config(
        page_title = page_title,
        page_icon = page_icon,
        layout="wide",
        initial_sidebar_state="expanded",
    )
    with open("assets/css/styles.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)